import io
import os
from docx import Document
from docx.shared import RGBColor, Inches
import markdown2
from bs4 import BeautifulSoup

# ==============================
# GENERADOR DE WORD (DOCX) - CORREGIDO
# ==============================

def process_rich_text(paragraph, html_content):
    """
    Procesa un párrafo HTML y añade 'runs' al párrafo de Word
    para manejar negritas (<strong>, <b>), cursivas (<em>, <i>) y código.
    """
    # Si recibimos un NavigableString (texto plano), lo agregamos directo
    if isinstance(html_content, str):
        paragraph.add_run(str(html_content))
        return

    # Si es un Tag, procesamos sus hijos
    for child in html_content.contents:
        if isinstance(child, str):
            paragraph.add_run(child)
        elif child.name in ['strong', 'b']:
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ['em', 'i']:
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            # Recursividad simple para otros tags anidados
            paragraph.add_run(child.get_text())

def generate_docx(markdown_text, title="Reporte Atelier", template_path=None):
    """
    Convierte Markdown a DOCX iterando secuencialmente para evitar duplicados.
    """
    try:
        # 1. Cargar Plantilla o Crear Nuevo
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            # Si no hay plantilla, agregamos el título manualmente
            doc.add_heading(title, 0)

        # 2. Convertir Markdown a HTML
        # Agregamos extras para asegurar estructura limpia
        html_content = markdown2.markdown(markdown_text, extras=["tables", "fenced-code-blocks"])
        soup = BeautifulSoup(html_content, "html.parser")

        # 3. Obtener contenedor principal (Body o Soup directo)
        container = soup.body if soup.body else soup

        # 4. ITERACIÓN SECUENCIAL (CORRECCIÓN ANTI-DUPLICADOS)
        # Iteramos solo los hijos directos. find_all era recursivo y causaba el error.
        for element in container.children:
            
            # Ignorar saltos de línea vacíos entre tags
            if isinstance(element, str):
                if element.strip(): 
                    doc.add_paragraph(element.strip())
                continue

            tag_name = element.name.lower()

            # --- ENCABEZADOS ---
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Extraer nivel del número (h1 -> 1, h2 -> 2)
                try:
                    level = int(tag_name[1])
                except: 
                    level = 1
                doc.add_heading(element.get_text().strip(), level=level)

            # --- PÁRRAFOS ---
            elif tag_name == 'p':
                p = doc.add_paragraph()
                process_rich_text(p, element)

            # --- LISTAS (UL / OL) ---
            elif tag_name in ['ul', 'ol']:
                # Iterar sobre los items de la lista <li>
                list_items = element.find_all('li', recursive=False)
                for li in list_items:
                    style = 'List Bullet' if tag_name == 'ul' else 'List Number'
                    try:
                        p = doc.add_paragraph(style=style)
                    except:
                        # Fallback si la plantilla no tiene el estilo
                        p = doc.add_paragraph(style='List Paragraph')
                        if tag_name == 'ul': p.style = 'List Bullet'
                    
                    process_rich_text(p, li)

            # --- CITAS (BLOCKQUOTE) ---
            elif tag_name == 'blockquote':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                # Las citas suelen tener <p> dentro, extraemos el texto o procesamos hijos
                text_content = element.get_text(" ", strip=True)
                run = p.add_run(text_content)
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)

            # --- CÓDIGO (PRE) ---
            elif tag_name == 'pre':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                run = p.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(9) if 'Pt' in globals() else None

            # --- OTROS (Fallback) ---
            else:
                text = element.get_text().strip()
                if text:
                    doc.add_paragraph(text)

        # 5. Guardar
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print(f"Error generando DOCX: {e}")
        return None
